# src/views/library_view.py
import streamlit as st
import pandas as pd
import os
import pickle
import json
import sys
import base64
import zipfile
import io
from io import StringIO
from docx import Document
from utils import state_management
from logic.data_processing import analyze_csv, expand_terms, run_analysis
from logic.models import OLSWrapper, SVRWrapper, RandomForestWrapper
from logic.hill_fit import detect_and_fit_hills
from logic.interpolation import interpolate_to_n_rows

def resource_path(relative_path):
    try:
        base_path = sys._MEIPASS
    except Exception:
        base_path = os.path.abspath(".")
    return os.path.join(base_path, relative_path)

PROJECT_LIBRARY_FILE = resource_path("project_library.json")
CONFIG_FILE = resource_path("config.json")

def load_library():
    if os.path.exists(PROJECT_LIBRARY_FILE):
        with open(PROJECT_LIBRARY_FILE, "r") as f:
            try:
                return json.load(f)
            except json.JSONDecodeError:
                return {}
    return {}

def save_library(library_data):
    temp_file = PROJECT_LIBRARY_FILE + ".tmp"
    with open(temp_file, "w") as f:
        json.dump(library_data, f, indent=4)
    if os.path.exists(PROJECT_LIBRARY_FILE):
        os.remove(PROJECT_LIBRARY_FILE)
    os.rename(temp_file, PROJECT_LIBRARY_FILE)

def load_config():
    if os.path.exists(CONFIG_FILE):
        with open(CONFIG_FILE, "r") as f:
            return json.load(f)
    return {"output_path": os.path.join(os.path.expanduser('~'), 'Desktop')}

def save_config(config):
    with open(CONFIG_FILE, "w") as f:
        json.dump(config, f)

def create_results_zip(project_name, analysis_run_name, analysis_run_data):
    zip_buffer = io.BytesIO()
    with zipfile.ZipFile(zip_buffer, "a", zipfile.ZIP_DEFLATED, False) as zip_file:
        model_type = analysis_run_data.get("model_type", "Unknown")
        for dep_var, results in analysis_run_data.get("models", {}).items():
            summary = results.get("summary_text", "No summary available.")
            doc = Document()
            doc.add_heading(f'Analysis for {project_name}', level=1)
            doc.add_heading(f'Run: {analysis_run_name} ({model_type})', level=2)
            doc.add_heading(f'Results for {dep_var}', level=3)
            doc.add_paragraph(summary)

            doc_buffer = io.BytesIO()
            doc.save(doc_buffer)
            doc_buffer.seek(0)
            zip_file.writestr(f"{project_name}_{analysis_run_name}_{dep_var}_Summary.docx", doc_buffer.read())

            csv_string = results.get("actual_vs_predicted_csv", "No data available.")
            zip_file.writestr(f"{project_name}_{analysis_run_name}_{dep_var}_ActualVsPredicted.csv", csv_string)

    zip_buffer.seek(0)
    return zip_buffer.getvalue()

def render():
    """Renders all UI components and logic for the Project Library tab."""
    st.subheader("🗂️ Project Library & Controls")

    with st.expander("Add New Dataset to Library", expanded=True):
        uploaded_file = st.file_uploader("Upload a CSV file to create a new project", type=["csv"])
        project_name_input = st.text_input("Enter a name for the new project")

        interpolate_enabled = st.checkbox(
            "Interpolate rows before saving",
            value=False,
            key="lib_interp_enable",
            help=(
                "Linearly interpolate the uploaded data to a larger row "
                "count before analysis. Use this when n_observations is "
                "less than the number of polynomial coefficients (1 + 2k + "
                "k(k-1)/2 for k independent variables), which is required "
                "for OLS / fitnlm to be identifiable. Mirrors the 13→15 "
                "pre-processing step in Ding et al., Adv. Therap. 2020."
            ),
        )
        interp_target_rows = None
        interp_time_col = None
        if interpolate_enabled:
            c_interp1, c_interp2 = st.columns([1, 2])
            interp_target_rows = c_interp1.number_input(
                "Target row count",
                min_value=2,
                value=15,
                step=1,
                key="lib_interp_n",
            )
            if uploaded_file is not None:
                try:
                    preview_df = pd.read_csv(uploaded_file)
                    uploaded_file.seek(0)
                    time_col_options = ["(use row index)"] + [
                        c for c in preview_df.columns
                        if pd.api.types.is_numeric_dtype(preview_df[c])
                    ]
                    chosen = c_interp2.selectbox(
                        "Time-axis column",
                        options=time_col_options,
                        key="lib_interp_time_col",
                        help="Column used as the x-axis for interpolation. "
                             "Leave on '(use row index)' if rows are already "
                             "evenly spaced in time.",
                    )
                    interp_time_col = None if chosen == "(use row index)" else chosen
                except Exception:
                    pass

        if st.button("Add to Library"):
            if uploaded_file is not None and project_name_input:
                library = load_library()
                if project_name_input in library:
                    st.error(f"A project named '{project_name_input}' already exists.")
                else:
                    df = pd.read_csv(uploaded_file)
                    if interpolate_enabled and interp_target_rows and interp_target_rows > len(df):
                        try:
                            df = interpolate_to_n_rows(
                                df,
                                n_target=int(interp_target_rows),
                                time_col=interp_time_col,
                            )
                            st.info(
                                f"Interpolated dataset from original rows to "
                                f"{len(df)} rows"
                                + (f" along '{interp_time_col}'." if interp_time_col else " along row index.")
                            )
                        except (ValueError, KeyError) as e:
                            st.error(f"Interpolation failed: {e}")
                            st.stop()
                    (data_df, _, independent_vars, dependent_vars,
                     _, _, _, _, _) = analyze_csv(df)
                    if not dependent_vars:
                        st.error(
                            "Validation Failed: No dependent variables found. "
                            "Please ensure dependent variable columns start with a recognized prefix "
                            "(e.g., 'Cell_', 'Patient_', 'Control_')."
                        )
                        st.stop()
                    if not independent_vars:
                        st.error(
                            "Validation Failed: No independent variables found. "
                            "Please check your column naming conventions."
                        )
                        st.stop()
                    df_json = df.to_json(orient='split')
                    library[project_name_input] = {
                        "data_df_json": df_json,
                        "analysis_runs": {} # Changed from analysis_results: None
                    }
                    save_library(library)
                    st.success(f"Successfully added project '{project_name_input}' to the library!")
                    st.rerun()
            else:
                st.warning("Please upload a file and provide a project name.")

    st.write("---")
    st.subheader("Existing Projects")

    library = load_library()
    projects = sorted(list(library.keys()))

    if not projects:
        st.info("No projects found. Use the section above to add your first project.")
        return

    selected_project = st.selectbox("Select a project", projects)

    # --- Project Actions ---
    col1, col2, col3 = st.columns(3)
    if col1.button("Load Project Data", type="primary"):
        project_data = library[selected_project]
        df = pd.read_json(StringIO(project_data['data_df_json']), orient='split')

        state_management.reset_state()
        (data_df, st.session_state.all_vars, st.session_state.independent_vars,
         st.session_state.dependent_vars, st.session_state.variable_stats,
         _, st.session_state.unique_variable_values,
         st.session_state.variable_descriptions,
         st.session_state.detected_binary_vars) = analyze_csv(df)
        st.session_state.exp_df = data_df
        st.session_state.processed_file = selected_project
        st.success(f"Loaded data for '{selected_project}'. You can now run an analysis.")
        st.rerun()

    if col2.button("Unload Project"):
        state_management.reset_state()
        st.success("Unloaded the current project.")
        st.rerun()

    if col3.button("Delete Project", type="secondary"):
        if selected_project in library:
            del library[selected_project]
            save_library(library)
        if st.session_state.get('processed_file') == selected_project:
            state_management.reset_state()
        st.success(f"Deleted project '{selected_project}'")
        st.rerun()

    # --- Analysis Management (only if a project's data is loaded) ---
    if st.session_state.get('processed_file') == selected_project:
        st.write("---")
        st.subheader(f"Analyses for: {selected_project}")
        
        project_data = library[selected_project]
        # Ensure 'analysis_runs' key exists
        if 'analysis_runs' not in project_data:
            project_data['analysis_runs'] = {}

        # --- Load Existing Analysis Section ---
        with st.container(border=True):
            st.markdown("##### Load Existing Analysis")
            analysis_runs = project_data.get("analysis_runs", {})
            if not analysis_runs:
                st.info("No analyses have been run for this project yet.")
            else:
                analysis_to_load = st.selectbox("Select analysis to load", options=list(analysis_runs.keys()))
                if st.button("Load Analysis"):
                    run_data = analysis_runs[analysis_to_load]
                    st.session_state.wrapped_models = {}
                    for dep_var, result_data in run_data.get("models", {}).items():
                        b64_key = "model_wrapper_b64"
                        if b64_key in result_data:
                            pickled_data = base64.b64decode(result_data[b64_key])
                            st.session_state.wrapped_models[dep_var] = pickle.loads(pickled_data)
                        else:
                            st.warning(f"Could not find a model for '{dep_var}' in '{analysis_to_load}'.")

                    expanded_df = st.session_state.exp_df.copy()
                    expand_terms(expanded_df, st.session_state.independent_vars)
                    st.session_state.expanded_df = expanded_df

                    # Hill fit detection (silent — never breaks load)
                    try:
                        st.session_state.hill_fits = detect_and_fit_hills(
                            st.session_state.exp_df,
                            st.session_state.independent_vars,
                            st.session_state.dependent_vars,
                        )
                    except Exception:
                        st.session_state.hill_fits = None

                    st.session_state.analysis_done = True
                    st.session_state.active_analysis_run = analysis_to_load # New state variable
                    st.success(f"Loaded analysis: '{analysis_to_load}'. Results are now active.")
                    st.rerun()

        # --- Run New Analysis Section ---
        with st.expander("Run a New Analysis", expanded=True):
            analysis_name = st.text_input("Enter a unique name for this new analysis", f"Analysis Run {len(analysis_runs) + 1}")
            model_type = st.selectbox("Select Regression Method", ["Polynomial OLS", "Nonlinear LS (fitnlm)", "Ridge Regression", "SVR", "Random Forest"])
            model_params = {}
            if model_type == 'Nonlinear LS (fitnlm)':
                st.info(
                    "Levenberg-Marquardt fit of the same quadratic polynomial as OLS "
                    "(Matlab `fitnlm`-equivalent). Unconstrained + unscaled: "
                    "mathematically identical to Polynomial OLS."
                )
                model_params['scale_inputs'] = st.checkbox(
                    "Standardize inputs (zero-mean, unit-variance)",
                    value=False,
                    key="fitnlm_scale_inputs",
                    help=(
                        "Recommended when independent variables span orders of "
                        "magnitude (e.g. ng/mL drug concentrations 10^1-10^6). "
                        "StandardScaler fits during training; predictions via "
                        "wrapper.predict() automatically apply the same transform, "
                        "so the user interacts with raw units throughout."
                    ),
                )
                model_params['log_transform'] = st.checkbox(
                    "Fit in log-dose space",
                    value=False,
                    key="fitnlm_log_transform",
                    help=(
                        "Apply log(x) to each independent variable before "
                        "fitting. Common pharmacology convention (Hill/"
                        "sigmoidal dose-response). Requires all inputs > 0."
                    ),
                )
            if model_type == 'Ridge Regression':
                st.info(
                    "L2-regularized polynomial fit. Handles rank-deficient data "
                    "(where OLS/fitnlm produce wild coefficients or refuse to "
                    "run) and suppresses edge-of-box extrapolation."
                )
                model_params['alpha'] = st.number_input(
                    "alpha (L2 regularization strength)",
                    min_value=1e-8,
                    max_value=1e4,
                    value=1e-3,
                    step=1e-3,
                    format="%.6f",
                    key="ridge_alpha",
                    help="0 -> OLS. Higher -> stronger shrinkage. Typical: 1e-3 to 1e+1.",
                )
                model_params['scale_inputs'] = st.checkbox(
                    "Standardize inputs (recommended, Ridge is scale-sensitive)",
                    value=True,
                    key="ridge_scale_inputs",
                )
                model_params['log_transform'] = st.checkbox(
                    "Fit in log-dose space",
                    value=False,
                    key="ridge_log_transform",
                    help=(
                        "Apply log(x) to each independent variable before "
                        "fitting. Common pharmacology convention. Requires "
                        "all inputs > 0."
                    ),
                )
            if model_type == 'SVR':
                st.info("SVR works best with scaled data and is sensitive to parameters.")
                model_params['C'] = st.number_input("C (Regularization)", value=1.0, min_value=0.01, step=0.1, format="%.2f")
                model_params['gamma'] = st.select_slider("Gamma", options=['scale', 'auto', 0.01, 0.1, 1.0, 10.0])
            elif model_type == 'Random Forest':
                st.info("Random Forest is a machine learning technique that works by building multiple decision trees to improve accuracy and reduce overfitting.")
                model_params['n_estimators'] = st.number_input("Number of Estimators (Trees)", value=100, min_value=10, step=10)
                model_params['max_depth'] = st.selectbox("Max Depth of Trees", options=[None, 5, 10, 20, 30], index=2, format_func=lambda x: "None" if x is None else str(x))

            if st.button("Run and Save Analysis", type="primary"):
                if not analysis_name:
                    st.error("Please enter a name for the analysis.")
                elif analysis_name in analysis_runs:
                    st.error(f"An analysis named '{analysis_name}' already exists for this project. Please choose a unique name.")
                else:
                    with st.spinner(f"Running {model_type} regression..."):
                        try:
                            # Use the original dataframe for non-polynomial models
                            df_for_analysis = st.session_state.exp_df.copy()
                            if model_type == 'Polynomial OLS':
                                expand_terms(df_for_analysis, st.session_state.independent_vars)
                            st.session_state.expanded_df = df_for_analysis # Store the df used for analysis
                            
                            models_to_store = {}
                            current_wrapped_models = {}

                            for dep_var in st.session_state.dependent_vars:
                                clean_df = df_for_analysis.dropna(subset=[dep_var] + st.session_state.independent_vars).copy()
                                wrapped_model = run_analysis(
                                    clean_df, st.session_state.independent_vars, dep_var, model_type, model_params
                                )
                                y_actual = clean_df[dep_var]
                                y_predicted = wrapped_model.predict(clean_df)
                                avp_df = pd.DataFrame({'Actual': y_actual, 'Predicted': y_predicted})
                                
                                current_wrapped_models[dep_var] = wrapped_model
                                pickled_model = pickle.dumps(wrapped_model)
                                b64_model = base64.b64encode(pickled_model).decode('utf-8')
                                
                                models_to_store[dep_var] = {
                                    "model_wrapper_b64": b64_model,
                                    "summary_text": wrapped_model.get_summary(),
                                    "actual_vs_predicted_csv": avp_df.to_csv(index=False)
                                }
                            
                            # Save to library
                            current_library = load_library()
                            current_library[selected_project]['analysis_runs'][analysis_name] = {
                                "model_type": model_type,
                                "models": models_to_store
                            }
                            save_library(current_library)
                            
                            # Activate the new analysis in the session
                            st.session_state.wrapped_models = current_wrapped_models

                            # Hill fit detection (silent — never breaks analysis)
                            try:
                                st.session_state.hill_fits = detect_and_fit_hills(
                                    st.session_state.exp_df,
                                    st.session_state.independent_vars,
                                    st.session_state.dependent_vars,
                                )
                            except Exception:
                                st.session_state.hill_fits = None

                            st.session_state.analysis_done = True
                            st.session_state.active_analysis_run = analysis_name
                            st.success(f"Analysis '{analysis_name}' complete and saved to project!")
                            st.rerun()

                        except Exception as e:
                            st.error(f"Analysis failed: {e}")
                            st.exception(e)
        
        # --- Download Results Section ---
        if st.session_state.get('analysis_done', False):
            st.write("---")
            active_run_name = st.session_state.get('active_analysis_run', 'N/A')
            st.subheader(f"⬇️ Download Results for Active Analysis: '{active_run_name}'")
            
            analysis_run_data = analysis_runs.get(active_run_name)
            if analysis_run_data:
                zip_data = create_results_zip(selected_project, active_run_name, analysis_run_data)
                st.download_button(
                    label=f"Download '{active_run_name}' Results (.zip)",
                    data=zip_data,
                    file_name=f"{selected_project}_{active_run_name}_Results.zip",
                    mime="application/zip",
                    type="primary"
                )
