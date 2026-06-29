# src/views/diagnostics_view.py

import streamlit as st
import pandas as pd
import numpy as np
import plotly.express as px
import plotly.graph_objects as go
from scipy import stats
import io
from docx import Document  # Required for Word report generation

from logic.diagnostics import (
    calculate_vif, 
    perform_normality_test, 
    perform_heteroscedasticity_test,
    perform_autocorrelation_test,
    perform_kfold_cv,
    perform_bootstrap_analysis,
    generate_diagnostics_report,
    generate_full_project_report
)
from logic.models import OLSWrapper

def apply_custom_layout(fig, height, width, title_size, axis_size, tick_size, 
                        show_grid=False, journal_style=False, legend_size=14,
                        title_text=None, x_text=None, y_text=None):
    """
    Applies consistent styling with High-Contrast BLACK text.
    """
    # 1. Base Layout (Force Black Fonts everywhere)
    update_dict = {
        'height': height,
        'width': width,
        'title': dict(
            font=dict(size=title_size, color='black'), # Force Black
            x=0.5,              
            xanchor='center',   
            yanchor='top'       
        ),
        'xaxis': dict(
            title_font=dict(size=axis_size, color='black'), # Force Black
            tickfont=dict(size=tick_size, color='black')    # Force Black
        ),
        'yaxis': dict(
            title_font=dict(size=axis_size, color='black'), # Force Black
            tickfont=dict(size=tick_size, color='black')    # Force Black
        ),
        'legend': dict(font=dict(size=legend_size, color='black')), # Force Black
        'margin': dict(l=80, r=40, t=80, b=80),
        'plot_bgcolor': 'white'
    }
    
    # 2. Journal Style Overrides
    if journal_style:
        axis_style = dict(
            showline=True,      
            linewidth=2,        
            linecolor='black',  
            mirror=True,        
            ticks='outside',    
            tickwidth=2,
            tickcolor='black',
            ticklen=6,
            showgrid=show_grid  
        )
        update_dict['xaxis'].update(axis_style)
        update_dict['yaxis'].update(axis_style)
    else:
        # Even in non-journal mode, keep text black but use default grid logic
        update_dict['xaxis']['showgrid'] = show_grid
        update_dict['yaxis']['showgrid'] = show_grid
    
    # 3. Apply Custom Text Labels
    if title_text is not None:
        update_dict['title']['text'] = title_text
    if x_text is not None:
        update_dict['xaxis']['title'] = dict(text=x_text, font=dict(size=axis_size, color='black'))
    if y_text is not None:
        update_dict['yaxis']['title'] = dict(text=y_text, font=dict(size=axis_size, color='black'))
        
    fig.update_layout(**update_dict)
    return fig

def create_word_report(model_wrapper, model_name, dataframe, independent_vars):
    """Generates a Word document with OLS diagnostics."""
    doc = Document()
    doc.add_heading(f'Diagnostic Report: {model_name}', 0)
    doc.add_paragraph(f"Generated on: {pd.Timestamp.now().strftime('%Y-%m-%d %H:%M:%S')}")

    # 1. Model Summary
    doc.add_heading('1. Model Summary', level=1)
    summary = model_wrapper.get_summary()
    doc.add_paragraph(summary) 

    # 2. VIF Analysis
    doc.add_heading('2. Multicollinearity (VIF)', level=1)
    try:
        # Re-calculate VIF
        effective_vars = independent_vars if independent_vars else model_wrapper.independent_vars
        vif_df = calculate_vif(model_wrapper, dataframe=dataframe, independent_vars=effective_vars)
        
        # --- FIX: ROBUST COLUMN HANDLING ---
        # If 'Variable' is not a column, it is likely in the index or named something else.
        if 'Variable' not in vif_df.columns:
            vif_df = vif_df.reset_index()
            # After reset, standardize the column names if we have exactly 2 columns (Name, Value)
            if len(vif_df.columns) == 2:
                vif_df.columns = ['Variable', 'VIF']
        # -----------------------------------
        
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Variable'
        hdr_cells[1].text = 'VIF Score'
        
        for _, row in vif_df.iterrows():
            row_cells = table.add_row().cells
            row_cells[0].text = str(row['Variable'])
            row_cells[1].text = f"{row['VIF']:.4f}"
    except Exception as e:
        doc.add_paragraph(f"Could not perform VIF analysis: {e}")

    # 3. Normality Test
    doc.add_heading('3. Normality of Residuals', level=1)
    try:
        residuals = model_wrapper.model.resid
        stat, p_val, test_name = perform_normality_test(residuals)
        
        p = doc.add_paragraph()
        p.add_run(f"Test Performed: ").bold = True
        p.add_run(f"{test_name}")
        
        p2 = doc.add_paragraph()
        p2.add_run(f"P-Value: ").bold = True
        p2.add_run(f"{p_val:.4f}")
        
        result_text = "Reject H0 (Residuals are likely NOT normal)" if p_val < 0.05 else "Fail to Reject H0 (Residuals appear normal)"
        p3 = doc.add_paragraph()
        p3.add_run("Conclusion: ").bold = True
        p3.add_run(result_text)
    except Exception as e:
        doc.add_paragraph(f"Could not perform normality test: {e}")

    # 4. Homoscedasticity
    doc.add_heading('4. Homoscedasticity', level=1)
    try:
        lm_p = perform_heteroscedasticity_test(model_wrapper.model.resid, model_wrapper)
        p = doc.add_paragraph()
        p.add_run("Breusch-Pagan p-value: ").bold = True
        p.add_run(f"{lm_p:.4f}")
        
        result_text = "Heteroscedasticity detected (Variance is not constant)" if lm_p < 0.05 else "Homoscedasticity indicated (Variance is constant)"
        p2 = doc.add_paragraph()
        p2.add_run("Conclusion: ").bold = True
        p2.add_run(result_text)
    except Exception as e:
        doc.add_paragraph(f"Could not perform heteroscedasticity test: {e}")
        
    # Save to buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def render():
    st.subheader("🔍 OLS Assumption & Uncertainty Diagnostics")
    st.info("Evaluate statistical assumptions (Normality, VIF) and predictive uncertainty.")

    if 'analysis_done' not in st.session_state or not st.session_state.analysis_done:
        st.warning("Please load a project and run an analysis first to generate models.")
        return

    # --- 1. GLOBAL STYLING & EXPORT SETTINGS ---
    with st.expander("🎨 Graph Appearance & Publication Settings", expanded=False):
        st.markdown("##### 📏 Dimensions & Quality")
        c1, c2, c3 = st.columns(3)
        plot_height = c1.number_input("Height (px)", 400, 2000, 600, step=50)
        plot_width = c2.number_input("Width (px)", 400, 3000, 800, step=50, help="Set Height = Width for a square plot.")
        export_scale = c3.selectbox("Export Scale (DPI)", [1, 2, 3, 4], index=2, help="3x = 300 DPI (Print Quality)")

        st.markdown("##### ✒️ Fonts & Style")
        c4, c5, c6 = st.columns(3)
        title_font_size = c4.number_input("Title Size", 10, 50, 25)
        axis_font_size = c5.number_input("Axis Label Size", 8, 40, 25)
        tick_font_size = c6.number_input("Tick Label Size", 8, 30, 25)
        
        c7, c8, c9 = st.columns(3)
        legend_font_size = c7.number_input("Legend Text Size", 8, 30, 25)
        journal_style = c8.checkbox("Journal Style (Boxed)", value=True, help="Adds a solid black frame, removes grey grid, and points ticks outward.")
        show_grid = c9.checkbox("Show Gridlines", value=False, help="Uncheck for clean white background.")

        download_config = {
            'toImageButtonOptions': {
                'format': 'png',
                'filename': 'high_res_plot',
                'height': plot_height,
                'width': plot_width,
                'scale': export_scale
            }
        }

    # Filter for OLS models only
    wrapped_models = st.session_state.get('wrapped_models', {})
    ols_models = {k: v for k, v in wrapped_models.items() if isinstance(v, OLSWrapper)}
    
    # Retrieve 'exp_df'
    data_df = st.session_state.get('exp_df', None) 
    independent_vars = st.session_state.get('independent_vars', [])
    
    if not ols_models:
        st.warning("⚠️ No OLS models found.")
        return

    st.divider()

    for model_name, model_wrapper in ols_models.items():
        # Safe key prefix (strip chars that break Streamlit keys)
        mk = model_name.replace(" ", "_").replace("(", "").replace(")", "").replace("/", "_")

        with st.expander(f"📊 {model_name}", expanded=True):

            # --- Report buttons ---
            col_btn1, col_btn2, _ = st.columns([1, 1, 2])
            with col_btn1:
                if st.button("📄 TXT Report", key=f"txt_btn_{mk}"):
                    with st.spinner("Generating text report..."):
                        report_text = generate_diagnostics_report(
                            model_wrapper,
                            model_name=model_name,
                            dataframe=data_df,
                            independent_vars=independent_vars
                        )
                        st.download_button(
                            label="Download TXT",
                            data=report_text,
                            file_name=f"diagnostics_report_{model_name}.txt",
                            mime="text/plain",
                            key=f"dl_txt_{mk}"
                        )
            with col_btn2:
                if st.button("📝 Word Report", key=f"word_btn_{mk}"):
                    with st.spinner("Generating Word report..."):
                        docx_buffer = create_word_report(
                            model_wrapper,
                            model_name,
                            data_df,
                            independent_vars
                        )
                        st.download_button(
                            label="Download DOCX",
                            data=docx_buffer,
                            file_name=f"diagnostics_report_{model_name}.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            key=f"dl_docx_{mk}"
                        )

            # Extract model data
            results = model_wrapper.model
            residuals = results.resid
            fitted_values = results.fittedvalues
            y_actual = results.model.endog

            # --- Tabbed Interface for Diagnostics ---
            tab1, tab2, tab3, tab4, tab5 = st.tabs([
                "1. Multicollinearity",
                "2. Normality",
                "3. Homoscedasticity",
                "4. Independence",
                "5. Predictive Uncertainty"
            ])

            # --- 1. Multicollinearity (VIF) ---
            with tab1:
                st.markdown("#### Variance Inflation Factor (VIF)")
                use_centered = st.checkbox("Apply Mean-Centering (Fix Structural Multicollinearity)", value=True, key=f"centered_{mk}")

                try:
                    effective_vars = independent_vars if independent_vars else model_wrapper.independent_vars

                    if use_centered:
                        if data_df is not None:
                            vif_df = calculate_vif(model_wrapper, dataframe=data_df, independent_vars=effective_vars)
                            st.success("✅ **Centered VIFs Active**")
                        else:
                            st.warning("⚠️ Original dataset missing. Showing Raw VIFs.")
                            vif_df = calculate_vif(model_wrapper)
                    else:
                        st.info("ℹ️ **Raw VIFs Active**")
                        vif_df = calculate_vif(model_wrapper)

                    def highlight_vif(val):
                        color = 'red' if val > 10 else ('orange' if val > 5 else 'green')
                        return f'color: {color}; font-weight: bold'

                    st.dataframe(vif_df.style.map(highlight_vif, subset=['VIF']).format({"VIF": "{:.2f}"}), width="stretch")
                except Exception as e:
                    st.error(f"Could not calculate VIF: {e}")

            # --- 2. Normality of Residuals ---
            with tab2:
                st.markdown("#### Normality of Residuals")
                col1, col2 = st.columns([1, 2])
                stat, p_val, test_name = perform_normality_test(residuals)

                with col1:
                    st.metric(f"{test_name} p-value", f"{p_val:.4f}")
                    if p_val < 0.05:
                        st.error("❌ **Reject H0**: Residuals NOT normal.")
                    else:
                        st.success("✅ **Fail to Reject H0**: Residuals look normal.")

                with col2:
                    with st.expander("✏️ Customize Q-Q Plot Labels"):
                        qq_title = st.text_input("Title", "Q-Q Plot", key=f"qq_title_{mk}")
                        qq_x = st.text_input("X Label", "Theoretical Quantiles", key=f"qq_x_{mk}")
                        qq_y = st.text_input("Y Label", "Ordered Values", key=f"qq_y_{mk}")

                    (osm, osr), (slope, intercept, r) = stats.probplot(residuals, dist="norm", plot=None)
                    fig_qq = go.Figure()
                    fig_qq.add_trace(go.Scatter(x=osm, y=osr, mode='markers', name='Residuals', marker=dict(color='black', symbol='circle-open', opacity=0.7, size=8)))
                    x_line = np.array([np.min(osm), np.max(osm)])
                    y_line = slope * x_line + intercept
                    fig_qq.add_trace(go.Scatter(x=x_line, y=y_line, mode='lines', name='Normal Line', line=dict(color='red', width=2)))

                    fig_qq = apply_custom_layout(
                        fig_qq, plot_height, plot_width, title_font_size, axis_font_size, tick_font_size,
                        show_grid, journal_style, legend_font_size, qq_title, qq_x, qq_y
                    )
                    st.plotly_chart(fig_qq, width="stretch", config=download_config)

                    st.divider()
                    with st.expander("✏️ Customize Histogram Labels"):
                        hist_title = st.text_input("Title", "Residual Histogram", key=f"hist_title_{mk}")
                        hist_x = st.text_input("X Label", "Residuals", key=f"hist_x_{mk}")
                        hist_y = st.text_input("Y Label", "Count", key=f"hist_y_{mk}")

                    fig_hist = px.histogram(x=residuals, nbins=30, title=hist_title, color_discrete_sequence=['lightgrey'])
                    fig_hist.update_traces(marker_line_color='black', marker_line_width=1.5, opacity=0.8)

                    fig_hist = apply_custom_layout(
                        fig_hist, plot_height, plot_width, title_font_size, axis_font_size, tick_font_size,
                        show_grid, journal_style, legend_font_size, hist_title, hist_x, hist_y
                    )
                    st.plotly_chart(fig_hist, width="stretch", config=download_config)

            # --- 3. Homoscedasticity ---
            with tab3:
                st.markdown("#### Homoscedasticity")
                col1, col2 = st.columns([1, 2])
                lm_p = perform_heteroscedasticity_test(residuals, model_wrapper)

                with col1:
                    st.metric("Breusch-Pagan p-value", f"{lm_p:.4f}")
                    if lm_p < 0.05:
                        st.error("❌ Heteroscedasticity detected.")
                    else:
                        st.success("✅ Variance is constant.")

                with col2:
                    with st.expander("✏️ Customize Plot Labels"):
                        rvf_title = st.text_input("Title", "Residuals vs. Fitted Values", key=f"rvf_title_{mk}")
                        rvf_x = st.text_input("X Label", "Fitted Values", key=f"rvf_x_{mk}")
                        rvf_y = st.text_input("Y Label", "Residuals", key=f"rvf_y_{mk}")

                    df_plot = pd.DataFrame({'Fitted': fitted_values, 'Residuals': residuals})
                    fig_rvf = px.scatter(df_plot, x='Fitted', y='Residuals', opacity=0.7)
                    fig_rvf.add_hline(y=0, line_dash="dash", line_color="red")
                    fig_rvf.update_traces(marker=dict(size=8, color='black', symbol='circle-open'))

                    fig_rvf = apply_custom_layout(
                        fig_rvf, plot_height, plot_width, title_font_size, axis_font_size, tick_font_size,
                        show_grid, journal_style, legend_font_size, rvf_title, rvf_x, rvf_y
                    )
                    st.plotly_chart(fig_rvf, width="stretch", config=download_config)

            # --- 4. Independence ---
            with tab4:
                st.markdown("#### Independence of Errors")
                dw_stat = perform_autocorrelation_test(residuals)
                st.metric("Durbin-Watson Statistic", f"{dw_stat:.4f}")

                with st.expander("✏️ Customize Plot Labels"):
                    ind_title = st.text_input("Title", "Residuals vs. Experiment Order", key=f"ind_title_{mk}")
                    ind_x = st.text_input("X Label", "Experiment Order (Index)", key=f"ind_x_{mk}")
                    ind_y = st.text_input("Y Label", "Residuals", key=f"ind_y_{mk}")

                fig_order = px.scatter(y=residuals)
                fig_order.add_hline(y=0, line_dash="dash", line_color="red")
                fig_order.update_traces(mode='lines+markers', marker=dict(color='black', size=6), line=dict(color='grey', width=1))

                fig_order = apply_custom_layout(
                    fig_order, plot_height, plot_width, title_font_size, axis_font_size, tick_font_size,
                    show_grid, journal_style, legend_font_size, ind_title, ind_x, ind_y
                )
                st.plotly_chart(fig_order, width="stretch", config=download_config)

                if 1.5 < dw_stat < 2.5:
                    st.success("✅ No Autocorrelation")
                else:
                    st.warning("⚠️ Possible Autocorrelation")

            # --- 5. Predictive Uncertainty ---
            with tab5:
                st.markdown("### Quantifying Uncertainty")
                st.markdown("#### Predicted vs. Observed")

                with st.expander("✏️ Customize Plot Labels", expanded=True):
                    pred_title = st.text_input("Title", "Predicted vs. Observed", key=f"pred_title_{mk}")
                    pred_x = st.text_input("X Label", "Observed (Actual)", key=f"pred_x_{mk}")
                    pred_y = st.text_input("Y Label", "Predicted (Model)", key=f"pred_y_{mk}")

                col_plot, _ = st.columns([2, 1])
                with col_plot:
                    min_val = min(min(y_actual), min(fitted_values))
                    max_val = max(max(y_actual), max(fitted_values))

                    fig_pred = go.Figure()
                    fig_pred.add_trace(go.Scatter(
                        x=y_actual, y=fitted_values, mode='markers',
                        name='Data', marker=dict(color='black', opacity=0.6, size=8, symbol='circle-open')
                    ))
                    fig_pred.add_trace(go.Scatter(
                        x=[min_val, max_val], y=[min_val, max_val],
                        mode='lines', name='Perfect Fit', line=dict(color='red', dash='dash')
                    ))

                    fig_pred = apply_custom_layout(
                        fig_pred, plot_height, plot_width, title_font_size, axis_font_size, tick_font_size,
                        show_grid, journal_style, legend_font_size, pred_title, pred_x, pred_y
                    )
                    st.plotly_chart(fig_pred, width="stretch", config=download_config)

                st.divider()

                col_cv, col_boot = st.columns(2)
                with col_cv:
                    with st.container(border=True):
                        st.markdown("#### 🔄 Cross-Validation (K-Fold)")
                        k_folds = st.number_input("Number of Folds (K)", min_value=2, max_value=20, value=5, key=f"kfolds_{mk}")
                        if st.button("Run K-Fold CV", key=f"run_cv_{mk}"):
                            try:
                                cv_res = perform_kfold_cv(model_wrapper, k=k_folds)
                                st.write("##### Results:")
                                st.metric("Avg RMSE", f"{cv_res['avg_rmse']:.4f}")
                                st.metric("Avg R²", f"{cv_res['avg_r2']:.4f}")
                            except Exception as e:
                                st.error(f"CV Failed: {e}")

                with col_boot:
                    with st.container(border=True):
                        st.markdown("#### 🎲 Bootstrap Analysis")
                        n_boot = st.number_input("Number of Resamples", min_value=10, max_value=1000, value=1000, key=f"nboot_{mk}")
                        if st.button("Run Bootstrap", key=f"run_boot_{mk}"):
                            try:
                                boot_df = perform_bootstrap_analysis(model_wrapper, n_bootstraps=n_boot)
                                st.dataframe(boot_df[['Term', 'Original', '95% CI Lower', '95% CI Upper', 'Stable?']].style.format({'Original': '{:.4f}', '95% CI Lower': '{:.4f}', '95% CI Upper': '{:.4f}'}), width="stretch", hide_index=True)
                            except Exception as e:
                                st.error(f"Bootstrap Failed: {e}")
