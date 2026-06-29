# src/logic/data_processing.py
import os
import pandas as pd
import numpy as np
import statsmodels.formula.api as smf
from docx import Document
import io
from datetime import datetime
from zipfile import ZipFile
from scipy.optimize import curve_fit
from .models import OLSWrapper, SVRWrapper, RandomForestWrapper, NonlinearLSWrapper, RidgeWrapper
from .helpers import _add_polynomial_terms, generate_prediction_grid


def _fit_hill(doses, effects):
    """
    Fit a 2-parameter Hill equation: E = d^h / (EC50^h + d^h).
    Returns (ec50, h) or (nan, nan) if fitting fails.
    Requires at least 3 non-zero dose points.
    """
    def hill_fn(d, ec50, h):
        return d**h / (ec50**h + d**h)

    valid = (doses > 1e-9) & np.isfinite(effects)
    d = doses[valid]
    e = np.clip(effects[valid], 1e-6, 1 - 1e-6)
    if len(d) < 3:
        return np.nan, np.nan
    try:
        popt, _ = curve_fit(
            hill_fn, d, e,
            p0=[np.median(d), 1.0],
            bounds=([1e-9, 0.1], [1e9, 10.0]),
            maxfev=2000
        )
        return float(popt[0]), float(popt[1])
    except Exception:
        return np.nan, np.nan


def _hill_predict(dose, ec50, h):
    """Predict effect at a given dose using fitted Hill parameters."""
    if np.isnan(ec50) or dose <= 1e-9:
        return 0.0
    return float(dose**h / (ec50**h + dose**h))

# ==========================================
#       PART 1: STANDARD ANALYSIS
# ==========================================

def analyze_csv(full_dataframe):
    """
    Inspects the input dataframe to prepare it for analysis.
    """
    first_row = full_dataframe.iloc[0]
    is_description_row = any(isinstance(item, str) for item in first_row)

    all_vars = full_dataframe.columns.tolist()

    if is_description_row:
        descriptions_row = first_row
        data_df = full_dataframe.iloc[1:].reset_index(drop=True)
        variable_descriptions = {var: descriptions_row[var] for var in all_vars}
    else:
        data_df = full_dataframe.copy()
        variable_descriptions = {var: var for var in all_vars}

    for col in data_df.columns:
        if data_df[col].dtype == 'object':
            try:
                cleaned_col = data_df[col].str.replace(',', '', regex=False)
                data_df[col] = pd.to_numeric(cleaned_col)
            except (AttributeError, ValueError, TypeError):
                pass

    ignore_prefixes = ["Cell_", "Zebrafish_", "Mouse_", "Rat_", "Patient_", "Control_"]
    independent_vars = [
        col for col in all_vars
        if col[0].isalpha() and not any(col.startswith(prefix) for prefix in ignore_prefixes)
    ]
    dependent_vars = [
        col for col in all_vars
        if any(col.startswith(prefix) for prefix in ignore_prefixes)
    ]

    variable_stats = {}
    unique_variable_values = {}
    detected_binary_vars = []

    for var in all_vars:
        numeric_col = pd.to_numeric(data_df[var], errors='coerce').dropna()
        if not numeric_col.empty:
            unique_sorted = np.sort(numeric_col.unique())
            unique_list = unique_sorted.tolist()
            unique_variable_values[var] = unique_list
            min_val = unique_sorted[0]
            max_val = unique_sorted[-1]
            second_min_val = unique_sorted[1] if len(unique_sorted) > 1 else min_val
            variable_stats[var] = (min_val, second_min_val, max_val)

            if var in independent_vars:
                if set(unique_list) == {0.0, 1.0} or set(unique_list) == {0, 1}:
                    detected_binary_vars.append(var)

    special_values_map = {}
    for col_name in all_vars:
        numeric_series = pd.to_numeric(data_df[col_name], errors='coerce')
        non_numeric_mask = numeric_series.isna() & data_df[col_name].notna()
        if non_numeric_mask.any():
            special_values = data_df[col_name][non_numeric_mask].unique().tolist()
            special_values_map[col_name] = special_values

    return (data_df, all_vars, independent_vars, dependent_vars,
            variable_stats, special_values_map, unique_variable_values,
            variable_descriptions, detected_binary_vars)

def expand_terms(dataframe, all_alphabet_vars):
    """A convenient wrapper for the polynomial term generation function."""
    _add_polynomial_terms(dataframe, all_alphabet_vars)

def generate_model_formula(C_col, all_alphabet_vars):
    """
    Constructs the model formula string required by the `statsmodels` library.
    """
    terms = []
    for header in all_alphabet_vars:
        terms.extend([header, f"{header}_sq"])

    if len(all_alphabet_vars) > 1:
        interaction_terms = [f"{var1}*{var2}" for i, var1 in enumerate(all_alphabet_vars) for var2 in all_alphabet_vars[i + 1:]]
        terms.extend(interaction_terms)

    return f"{C_col} ~ " + " + ".join(terms)

def run_analysis(dataframe, independent_vars, dependent_var, model_type, model_params={}, variable_descriptions=None):
    """
    Main analysis engine. Fits the specified model and returns a wrapped model object.
    """
    if model_type == 'Polynomial OLS':
        model_formula = generate_model_formula(dependent_var, independent_vars)
        ols_model = smf.ols(model_formula, data=dataframe).fit()
        return OLSWrapper(ols_model, model_formula)

    elif model_type == 'Nonlinear LS (fitnlm)':
        # Matlab fitnlm-equivalent on the same quadratic PRS polynomial.
        # Unconstrained + unscaled: numerically identical to Polynomial OLS.
        # model_params keys: 'bounds' (Lower/Upper for fitnlm), 'scale_inputs'
        # (standardize features; recommended when raw inputs span orders of
        # magnitude, e.g. ng/mL drug concentrations 10^1–10^6).
        nls_wrapper = NonlinearLSWrapper(independent_vars)
        nls_wrapper.fit(
            dataframe,
            dependent_var,
            bounds=model_params.get('bounds'),
            scale_inputs=model_params.get('scale_inputs', False),
            log_transform=model_params.get('log_transform', False),
        )
        return nls_wrapper

    elif model_type == 'Ridge Regression':
        # L2-regularized polynomial fit. Handles rank-deficient design
        # matrices (where OLS/fitnlm are under-identified or refuse to run)
        # by penalizing coefficient norms, which also suppresses edge-of-box
        # extrapolation in the response surface.
        ridge_wrapper = RidgeWrapper(independent_vars)
        ridge_wrapper.fit(
            dataframe,
            dependent_var,
            alpha=model_params.get('alpha', 1e-3),
            scale_inputs=model_params.get('scale_inputs', True),
            log_transform=model_params.get('log_transform', False),
        )
        return ridge_wrapper

    elif model_type == 'SVR':
        svr_wrapper = SVRWrapper(independent_vars)
        svr_wrapper.fit(dataframe, dependent_var, **model_params)
        return svr_wrapper

    elif model_type == 'Random Forest':
        rf_wrapper = RandomForestWrapper(independent_vars)
        model_params['variable_descriptions'] = variable_descriptions
        rf_wrapper.fit(dataframe, dependent_var, **model_params)
        return rf_wrapper
        
    else:
        raise ValueError(f"Unknown model type: {model_type}")

def predict_surface(model, all_alphabet_vars, x_var, y_var, fixed_vars_dict, x_range, y_range):
    """
    Generates the Z-axis surface data for 3D plotting.
    Required by plotting_view.py.
    """
    x_grid, y_grid, predict_df = generate_prediction_grid(model, x_var, y_var, fixed_vars_dict, x_range, y_range)
    predicted_z = model.predict(predict_df)

    if isinstance(predicted_z, pd.Series):
        predicted_z = predicted_z.values

    z_grid = predicted_z.reshape(x_grid.shape)
    return x_grid, y_grid, z_grid

def generate_optimization_report(report_data, variable_descriptions):
    """
    Generates a professional .docx report from optimization results.
    Required by optimizer_view.py.
    """
    doc = Document()
    doc.add_heading('Optimization Run Report', 0)

    doc.add_paragraph(f"Report Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph(f"Optimization Type: {report_data.get('optimization_type', 'N/A')}")

    doc.add_heading('Model Information', level=1)
    doc.add_paragraph(f"Model Name(s): {report_data.get('model_name', 'N/A')}")
    formula = report_data.get('model_formula', 'N/A')
    doc.add_paragraph("Model Formula:")
    doc.add_paragraph(str(formula), style='Intense Quote')

    doc.add_heading('Results', level=1)
    results = report_data.get('results', {})
    doc.add_paragraph(f"Status: {results.get('Status', 'N/A')}")

    if 'Top Solutions' in results and isinstance(results['Top Solutions'], list):
        for result_item in results['Top Solutions']:
            doc.add_heading(f"Rank {result_item['Rank']} Solution", level=2)
            doc.add_paragraph(f"Objective Outcome: {result_item['Final Objective Outcome']:.4f}")
            doc.add_paragraph("Optimal Dosages:")
            dosages = result_item.get('Optimal Dosages', {})
            if dosages:
                table = doc.add_table(rows=1, cols=2, style='Table Grid')
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = 'Variable'
                hdr_cells[1].text = 'Optimal Value'
                for var, val in dosages.items():
                    row_cells = table.add_row().cells
                    row_cells[0].text = variable_descriptions.get(var, var)
                    row_cells[1].text = f"{val:.4f}"
    
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream.getvalue()

def generate_combined_predictions_csv(wrapped_models, data_df, independent_vars):
    """
    Generates a CSV of actual vs predicted values.
    Required by diagnostics_view.py.
    """
    all_predictions_df = pd.DataFrame(index=data_df.index)

    for model_name, model_obj in wrapped_models.items():
        clean_df = data_df.dropna(subset=[model_name] + independent_vars).copy()
        y_actual = clean_df[model_name]
        y_predicted = model_obj.predict(clean_df)

        all_predictions_df[f'Actual_{model_name}'] = y_actual
        all_predictions_df[f'Predicted_{model_name}'] = pd.Series(y_predicted, index=y_actual.index)

    csv_buffer = io.StringIO()
    all_predictions_df.to_csv(csv_buffer, index_label='Experiment_Index')
    return csv_buffer.getvalue().encode('utf-8')

def generate_combined_report(report_data_list, variable_descriptions):
    doc = Document()
    doc.add_heading('Combined Analysis Report', 0)
    doc.add_paragraph(f"Report Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

    for report_data in report_data_list:
        model_name = report_data.get('model_name', 'N/A')
        doc.add_heading(f"Results for: {model_name}", level=1)
        summary_text = report_data.get('model_summary', 'Summary not available.')
        doc.add_paragraph(summary_text)

    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream.getvalue()

# ==========================================
#       PART 2: LEGACY SYNERGY (MATRIX)
# ==========================================

def calculate_synergy(dataframe, drug1_name, drug2_name, effect_name, model='gamma'):
    """
    Legacy Matrix-based synergy calculation.
    Kept to prevent breaking older imports if they exist.
    """
    df_clean = dataframe.copy()
    df_clean[drug1_name] = pd.to_numeric(df_clean[drug1_name], errors='coerce')
    df_clean[drug2_name] = pd.to_numeric(df_clean[drug2_name], errors='coerce')

    response_matrix = df_clean.pivot_table(
        index=drug1_name, columns=drug2_name, values=effect_name
    ).sort_index(axis=0).sort_index(axis=1)

    row_indexes = response_matrix.index.values
    col_indexes = response_matrix.columns.values
    
    # Find closest to 0
    row_0_candidates = row_indexes[np.abs(row_indexes) < 1e-6]
    col_0_candidates = col_indexes[np.abs(col_indexes) < 1e-6]

    if len(row_0_candidates) == 0 or len(col_0_candidates) == 0:
        return response_matrix # Fail gracefully-ish

    zero_row_val = row_0_candidates[0]
    zero_col_val = col_0_candidates[0]

    drug1_single = response_matrix[zero_col_val] 
    drug2_single = response_matrix.loc[zero_row_val]

    d1_grid = pd.DataFrame(
        np.tile(drug1_single.values[:, None], (1, response_matrix.shape[1])), 
        index=response_matrix.index, columns=response_matrix.columns
    )
    d2_grid = pd.DataFrame(
        np.tile(drug2_single.values[None, :], (response_matrix.shape[0], 1)), 
        index=response_matrix.index, columns=response_matrix.columns
    )

    if model.lower() == 'hsa':
        expected_matrix = np.maximum(d1_grid, d2_grid)
        synergy_matrix = response_matrix - expected_matrix
        return synergy_matrix
    else:
        # Default Gamma
        V_obs = 1 - response_matrix
        V_d1 = 1 - d1_grid
        V_d2 = 1 - d2_grid
        V_expected = np.minimum(V_d1, V_d2)
        epsilon = 1e-9 
        synergy_matrix = V_obs / (V_expected + epsilon)
        return synergy_matrix

# ==========================================
#       PART 3: NEW HIGH-THROUGHPUT SYNERGY
# ==========================================

def detect_effect_direction(df, dose_cols, effect_col):
    """
    Automatically detects if the 'effect' is Inhibition (Higher=Stronger) 
    or Viability (Lower=Stronger) based on correlation.
    """
    # Calculate simple correlation between sum of doses and effect
    total_dose = df[dose_cols].sum(axis=1)
    corr = total_dose.corr(df[effect_col])
    
    # If correlation is Positive (>0), Higher Dose = Higher Value -> Inhibition
    # If correlation is Negative (<0), Higher Dose = Lower Value -> Viability
    if corr < 0:
        return 'viability'
    else:
        return 'inhibition'

def get_single_agent_response(row_idx, current_row, drug_name, df_clean, dose_cols, effect_col, model=None):
    """
    Finds the single-agent effect for 'drug_name' at the dose specified in 'current_row'.
    1. Looks for exact experimental match in df_clean.
    2. If missing, uses 'model' to predict it.
    """
    target_dose = current_row[drug_name]
    
    # If dose is effectively 0, effect is 0 (relative to baseline)
    if target_dose < 1e-6:
        return 0.0

    # 1. Search for experimental data
    # Logic: Find rows where 'drug_name' is close to target_dose AND all other drugs are 0
    mask = (np.abs(df_clean[drug_name] - target_dose) < 1e-6)
    for other_drug in dose_cols:
        if other_drug != drug_name:
            mask &= (np.abs(df_clean[other_drug]) < 1e-6)
    
    matches = df_clean[mask]
    
    if not matches.empty:
        # Return mean of replicates
        return matches[effect_col].mean()
    
    # 2. Fallback: AI Prediction
    if model is not None:
        # Construct a synthetic row for prediction
        pred_input = {col: 0.0 for col in dose_cols}
        pred_input[drug_name] = target_dose
        pred_df = pd.DataFrame([pred_input])
        
        # Predict
        try:
            prediction = model.predict(pred_df)[0]
            # Clip prediction to reasonable bounds (0-1 for inhibition)
            return np.clip(prediction, 0.0, 1.0) 
        except:
            return np.nan
            
    return np.nan

def calculate_high_throughput_synergy(df, dose_cols, effect_col, model=None):
    """
    Row-based synergy calculation supporting N-drugs, missing data (AI-fill), 
    and Bliss/HSA scores.
    """
    results = df.copy()
    
    # 1. Auto-Detect Direction & Standardize to Inhibition (0=No Effect, 1=Full Kill)
    direction = detect_effect_direction(df, dose_cols, effect_col)
    
    # Create an internal working column 'std_effect' (Standardized Effect)
    if direction == 'viability':
        # Assume 0-1 range. Transform: Inhibition = 1 - Viability
        # Check max value to be safe
        max_val = df[effect_col].max()
        if max_val > 1.5: # Likely percentage 0-100
             results['std_effect'] = 1 - (df[effect_col] / 100.0)
        else:
             results['std_effect'] = 1 - df[effect_col]
    else:
        # Already inhibition. Check bounds.
        if df[effect_col].max() > 1.5:
            results['std_effect'] = df[effect_col] / 100.0
        else:
            results['std_effect'] = df[effect_col]
            
    # Lists to store results
    hsa_synergy_scores = []
    bliss_synergy_scores = []
    zip_synergy_scores = []
    notes = []

    # Pre-clean for lookups
    df_lookup = results.copy()
    for col in dose_cols:
        df_lookup[col] = pd.to_numeric(df_lookup[col], errors='coerce').fillna(0)

    # --- Pre-fit Hill curves for ZIP (requires model) ---
    # ZIP differs from Bliss: it uses curve-fitted single-agent effects (from Hill equation)
    # rather than single observed data points, making it robust to sparse/noisy reference data.
    effect_max_val = df[effect_col].max()
    drug_hill_params = {}
    if model is not None:
        for drug in dose_cols:
            drug_doses = np.sort(df_lookup[drug].unique())
            drug_doses = drug_doses[drug_doses > 1e-9]
            if len(drug_doses) < 3:
                drug_hill_params[drug] = (np.nan, np.nan)
                continue
            hill_effects = []
            for d in drug_doses:
                pred_input = {col: 0.0 for col in dose_cols}
                pred_input[drug] = float(d)
                try:
                    raw_eff = float(model.predict(pd.DataFrame([pred_input]))[0])
                    # Apply same direction transform used for std_effect
                    if direction == 'viability':
                        eff = 1 - (raw_eff / 100.0) if effect_max_val > 1.5 else 1 - raw_eff
                    else:
                        eff = raw_eff / 100.0 if effect_max_val > 1.5 else raw_eff
                    hill_effects.append(np.clip(eff, 0.0, 1.0))
                except Exception:
                    hill_effects.append(np.nan)
            drug_hill_params[drug] = _fit_hill(drug_doses, np.array(hill_effects))

    # Pre-compute single-agent responses into an O(1) lookup dict
    single_agent_cache = {}
    for drug in dose_cols:
        for _, cache_row in df_lookup.iterrows():
            dose = cache_row[drug]
            if dose < 1e-6:
                continue
            key = (drug, round(dose, 9))
            if key in single_agent_cache:
                continue
            mask = np.abs(df_lookup[drug] - dose) < 1e-6
            for other in dose_cols:
                if other != drug:
                    mask &= np.abs(df_lookup[other]) < 1e-6
            matches = df_lookup[mask]
            single_agent_cache[key] = matches['std_effect'].mean() if not matches.empty else np.nan

    # 2. Iterate Rows
    for idx, row in df_lookup.iterrows():
        # Identify Active Drugs in this row (Dose > 0)
        active_drugs = [d for d in dose_cols if row[d] > 1e-6]

        if len(active_drugs) < 2:
            # Not a combination
            hsa_synergy_scores.append(0.0)
            bliss_synergy_scores.append(0.0)
            zip_synergy_scores.append(np.nan)
            notes.append("Single Agent / Control")
            continue

        # Get Single Agent Effects (Observed or Predicted)
        single_agent_effects = []
        is_predicted = False

        for drug in active_drugs:
            target_dose = row[drug]
            if target_dose < 1e-6:
                eff = 0.0
            else:
                key = (drug, round(target_dose, 9))
                eff = single_agent_cache.get(key, np.nan)
                # Fallback: AI Prediction
                if pd.isna(eff) and model is not None:
                    pred_input = {col: 0.0 for col in dose_cols}
                    pred_input[drug] = target_dose
                    pred_df = pd.DataFrame([pred_input])
                    try:
                        prediction = model.predict(pred_df)[0]
                        eff = np.clip(prediction, 0.0, 1.0)
                    except Exception:
                        eff = np.nan
            if pd.isna(eff):
                is_predicted = True
            single_agent_effects.append(eff)
            
        if any(pd.isna(x) for x in single_agent_effects):
            hsa_synergy_scores.append(np.nan)
            bliss_synergy_scores.append(np.nan)
            zip_synergy_scores.append(np.nan)
            notes.append("Missing Reference Data")
            continue
            
        # 3. Calculate HSA (Highest Single Agent)
        # Exp_HSA = Max(Single Agent Effects)
        # Synergy = Obs - Exp
        obs_effect = row['std_effect']
        exp_hsa = max(single_agent_effects)
        hsa_score = obs_effect - exp_hsa
        
        # 4. Calculate Bliss (Independence)
        # Exp_Bliss = 1 - (product of (1-Effects))
        unaffected_product = 1.0
        for eff in single_agent_effects:
            unaffected_product *= (1.0 - eff)
        
        exp_bliss = 1.0 - unaffected_product
        bliss_score = obs_effect - exp_bliss
        
        # 5. ZIP (Zero Interaction Potency)
        # Uses Hill-curve-fitted single-agent effects instead of observed point values.
        # Expected ZIP = Bliss independence using curve-fitted E_A and E_B.
        # This makes ZIP robust to sparse/noisy single-agent reference data.
        zip_score = np.nan
        if drug_hill_params:
            hill_effects_for_row = []
            hill_valid = True
            for drug in active_drugs:
                ec50, h = drug_hill_params.get(drug, (np.nan, np.nan))
                if np.isnan(ec50):
                    hill_valid = False
                    break
                hill_effects_for_row.append(_hill_predict(row[drug], ec50, h))
            if hill_valid and len(hill_effects_for_row) >= 2:
                unaffected_zip = 1.0
                for e in hill_effects_for_row:
                    unaffected_zip *= (1.0 - e)
                exp_zip = 1.0 - unaffected_zip
                zip_score = obs_effect - exp_zip

        hsa_synergy_scores.append(hsa_score)
        bliss_synergy_scores.append(bliss_score)
        zip_synergy_scores.append(zip_score)
        notes.append("Calculated" if not is_predicted else "AI Estimated Ref")

    # 6. Append raw scores
    results['HSA_Synergy'] = hsa_synergy_scores
    results['Bliss_Synergy'] = bliss_synergy_scores
    results['ZIP_Synergy'] = zip_synergy_scores
    results['Analysis_Note'] = notes

    # 7. Z-scores (combination rows only; single-agent rows stay NaN)
    combo_mask = results['Analysis_Note'] != 'Single Agent / Control'
    for col in ['HSA_Synergy', 'Bliss_Synergy', 'ZIP_Synergy']:
        z_col = f'{col}_Z'
        results[z_col] = np.nan
        valid = results.loc[combo_mask, col].dropna()
        if len(valid) > 1:
            mean_s, std_s = valid.mean(), valid.std()
            if std_s > 1e-9:
                results.loc[combo_mask, z_col] = (results.loc[combo_mask, col] - mean_s) / std_s

    return results, direction
